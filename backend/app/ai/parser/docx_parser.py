"""
DOCX document parser.
Uses python-docx to extract text, paragraphs, headers, and tables.
"""

import docx
import structlog

logger = structlog.get_logger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract raw text from a Microsoft Word (.docx) document.
    Includes text from paragraphs and tables.
    """
    text_chunks = []
    try:
        doc = docx.Document(file_path)

        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_chunks.append(paragraph.text.strip())

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))

        full_text = "\n".join(text_chunks)
        logger.info("DOCX text extracted successfully", char_count=len(full_text))
        return full_text
    except Exception as e:
        logger.error("Failed to extract text from DOCX file", path=file_path, error=str(e))
        return ""
